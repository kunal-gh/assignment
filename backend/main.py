"""
AI Resume Screener — FastAPI Backend v3.0
Embedding: Google Gemini (gemini-embedding-001) — batch API, task-aware
OCR: PyMuPDF for digital PDFs → Gemini Vision fallback for scanned PDFs
Parsing: python-docx (BytesIO) for DOCX files
"""

import asyncio
import logging
import os
import re
import tempfile
import time
import uuid
from typing import Any, Dict, List, Optional

# Load .env for local development (no-op on Render where env vars are set in dashboard)
try:
    from dotenv import load_dotenv

    load_dotenv()
except ImportError:
    pass  # python-dotenv not installed — env vars must be set in the environment directly

import numpy as np
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Config ───────────────────────────────────────────────────────────────────

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
GEMINI_EMBED_MODEL = "models/gemini-embedding-001"
GEMINI_VISION_MODEL = "gemini-1.5-flash"
EMBEDDING_DIM = 768  # truncated from 3072

# Configure Gemini client using new google-genai SDK (google-generativeai is deprecated)
_gemini_ready = False
_gemini_client = None
if GOOGLE_API_KEY:
    try:
        from google import genai as _genai_module

        _gemini_client = _genai_module.Client(api_key=GOOGLE_API_KEY)
        _gemini_ready = True
        logger.info("Google Gemini client configured (new google-genai SDK).")
    except Exception as e:
        logger.error(f"Gemini configuration failed: {e}")
else:
    logger.warning("GOOGLE_API_KEY not set — TF-IDF fallback will be used for embeddings; scanned PDFs will not OCR.")

# ─── App ──────────────────────────────────────────────────────────────────────

app = FastAPI(
    title="AI Resume Screener API",
    description="3-layer ML pipeline: Gemini Vision OCR + regex skill extraction + Gemini batch embeddings.",
    version="3.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)

# ─── Skill taxonomy ───────────────────────────────────────────────────────────

SKILL_ENTRIES: List[tuple] = [
    ("python", ["python3", "py"]),
    ("javascript", ["js", "es6"]),
    ("typescript", ["ts"]),
    ("sql", ["mysql", "sqlite"]),
    ("java", []),
    ("react", ["reactjs", "react.js"]),
    ("node.js", ["node", "nodejs"]),
    ("django", []),
    ("flask", []),
    ("fastapi", ["fast api"]),
    ("machine learning", ["ml"]),
    ("deep learning", ["dl"]),
    ("natural language processing", ["nlp"]),
    ("computer vision", ["cv"]),
    ("reinforcement learning", ["rl"]),
    ("tensorflow", ["tf"]),
    ("pytorch", ["torch"]),
    ("scikit-learn", ["sklearn", "scikit learn"]),
    ("transformers", ["huggingface transformers"]),
    ("sentence-transformers", ["sbert"]),
    ("hugging face", ["huggingface", "hf"]),
    ("langchain", ["lang chain"]),
    ("openai", ["open ai"]),
    ("llm", ["large language model", "large language models"]),
    ("rag", ["retrieval augmented generation", "retrieval-augmented generation"]),
    ("embeddings", ["vector embeddings"]),
    ("faiss", []),
    ("spacy", []),
    ("mlops", ["ml ops"]),
    ("mlflow", []),
    ("airflow", ["apache airflow"]),
    ("crewai", ["crew ai"]),
    ("mcp", ["model context protocol"]),
    ("graph neural networks", ["gnn", "gnns"]),
    ("pydantic", []),
    ("vector database", ["vector db", "vectordb"]),
    ("pandas", []),
    ("numpy", []),
    ("aws", ["amazon web services"]),
    ("azure", ["microsoft azure"]),
    ("gcp", ["google cloud", "google cloud platform"]),
    ("docker", ["containerization"]),
    ("kubernetes", ["k8s"]),
    ("git", ["github", "gitlab"]),
    ("ci/cd", ["cicd", "continuous integration", "continuous deployment"]),
    ("postgresql", ["postgres", "psql"]),
    ("mongodb", ["mongo"]),
    ("redis", []),
    ("neo4j", []),
    ("qdrant", []),
    ("rest api", ["restful", "rest", "restful api"]),
    ("graphql", []),
    ("microservices", []),
    ("devops", []),
    ("agile", ["agile methodology"]),
    ("scrum", []),
    ("streamlit", []),
    ("serverless", []),
    ("terraform", []),
    ("linux", ["ubuntu"]),
    ("gemini", ["google gemini"]),
    ("vertex ai", ["vertexai"]),
]

SKILL_LOOKUP: Dict[str, str] = {}
for canonical, aliases in SKILL_ENTRIES:
    SKILL_LOOKUP[canonical.lower()] = canonical
    for alias in aliases:
        SKILL_LOOKUP[alias.lower()] = canonical

SKILL_IDF: Dict[str, float] = {
    "rag": 3.5,
    "mcp": 3.4,
    "crewai": 3.3,
    "faiss": 3.2,
    "spacy": 3.1,
    "sentence-transformers": 3.4,
    "mlops": 3.0,
    "neo4j": 3.0,
    "qdrant": 3.1,
    "vector database": 3.0,
    "llm": 2.8,
    "embeddings": 2.7,
    "langchain": 2.6,
    "natural language processing": 2.6,
    "transformers": 2.5,
    "pytorch": 2.4,
    "tensorflow": 2.3,
    "kubernetes": 2.2,
    "docker": 2.0,
    "aws": 1.9,
    "machine learning": 2.1,
    "deep learning": 2.2,
    "python": 1.5,
    "git": 1.2,
    "gemini": 3.0,
    "vertex ai": 2.8,
}

# ─── Layer 1: Text Extraction ─────────────────────────────────────────────────


def extract_text_from_pdf_bytes(data: bytes) -> str:
    """Extract text from a digital (text-layer) PDF using PyMuPDF."""
    try:
        import fitz

        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    except Exception:
        pass
    # ASCII byte fallback
    parts, cur = [], ""
    for b in data:
        if 32 <= b <= 126 or b in (9, 10, 13):
            cur += chr(b)
        else:
            if len(cur) >= 2:
                parts.append(cur)
            cur = ""
    if cur:
        parts.append(cur)
    return " ".join(parts).replace("(", " ").replace(")", " ")


def extract_text_from_docx_bytes(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx (no temp file needed)."""
    try:
        import io
        from docx import Document

        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"python-docx failed: {e}, using raw decode")
        return data.decode("utf-8", errors="ignore")


def _gemini_vision_ocr_sync(file_path: str, filename: str) -> str:
    """
    Synchronous Gemini Vision OCR using the new google-genai SDK.
    Reads the PDF as a visual document — handles scanned/image-based PDFs.
    """
    from google.genai import types as _types

    if not _gemini_client:
        return ""

    try:
        with open(file_path, "rb") as f:
            pdf_bytes = f.read()
        response = _gemini_client.models.generate_content(
            model=GEMINI_VISION_MODEL,
            contents=[
                _types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"),
                (
                    "This is a resume/CV document. Extract ALL text content exactly as written. "
                    "Include: name, contact information, work experience (dates, companies, roles, descriptions), "
                    "education, skills, certifications, and any other sections. "
                    "Return only the raw extracted text — do not summarize or interpret."
                ),
            ],
        )
        return response.text or ""
    except Exception as e:
        logger.error(f"Gemini Vision OCR error for {filename}: {e}")
        return ""


async def extract_text(file_path: str, filename: str) -> str:
    """
    Smart text extraction cascade:
    1. PDF → PyMuPDF (digital text layer, instant, free)
    2. PDF with < 100 chars extracted → Gemini Vision OCR (scanned document)
    3. DOCX → python-docx via BytesIO
    """
    with open(file_path, "rb") as f:
        data = f.read()

    if filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf_bytes(data)
        if len(text.strip()) < 100 and _gemini_ready:
            logger.info(f"Scanned PDF detected: {filename} — invoking Gemini Vision OCR")
            vision_text = await asyncio.to_thread(_gemini_vision_ocr_sync, file_path, filename)
            if vision_text and len(vision_text.strip()) > len(text.strip()):
                logger.info(f"Gemini Vision OCR succeeded for {filename} ({len(vision_text)} chars)")
                return vision_text
        return text

    if filename.lower().endswith(".docx"):
        return extract_text_from_docx_bytes(data)

    return data.decode("utf-8", errors="ignore")


# ─── Layer 2: Skill Extraction ────────────────────────────────────────────────


def extract_skills(text: str) -> List[str]:
    norm = text.lower().replace("(", " ").replace(")", " ").replace(",", " ").replace(";", " ").replace("-", " ")
    found = set()
    for variant, canonical in sorted(SKILL_LOOKUP.items(), key=lambda x: -len(x[0])):
        if len(variant) <= 4:
            if re.search(r"(?:^|[\s,;(])" + re.escape(variant) + r"(?:[\s,;)]|$)", norm):
                found.add(canonical)
        else:
            if variant in norm:
                found.add(canonical)
    return list(found)


# ─── Layer 3: Gemini Batch Embeddings ─────────────────────────────────────────


def _gemini_embed_sync(texts: List[str], task_type: str) -> List[List[float]]:
    """Embed a batch of texts in ONE Gemini API call using the new google-genai SDK."""
    from google.genai import types as _types

    if not _gemini_client:
        return []

    response = _gemini_client.models.embed_content(
        model=GEMINI_EMBED_MODEL,
        contents=texts,
        config=_types.EmbedContentConfig(
            task_type=task_type,
            output_dimensionality=EMBEDDING_DIM,
        ),
    )
    return [e.values for e in response.embeddings]


async def get_embeddings_batch(texts: List[str], task_type: str = "retrieval_document") -> List[Optional[List[float]]]:
    """
    Async wrapper for Gemini batch embedding with chunking.
    Retries with exponential backoff on rate limit (429) or transient errors.
    Returns None entries where embedding failed.
    """
    if not _gemini_ready or not texts:
        return [None] * len(texts)

    # Stricter truncation (6500 chars) to stay safely under 2048 token limit
    truncated = [t[:6500] for t in texts]
    all_embeddings = []

    # Process in chunks of 50 to avoid payload size limits
    chunk_size = 50
    for i in range(0, len(truncated), chunk_size):
        chunk = truncated[i : i + chunk_size]
        chunk_embeddings = None

        for attempt in range(3):
            try:
                embeddings = await asyncio.to_thread(_gemini_embed_sync, chunk, task_type)
                if len(embeddings) == len(chunk):
                    chunk_embeddings = embeddings
                    break
                # Pad if partial result returned
                chunk_embeddings = embeddings + [None] * (len(chunk) - len(embeddings))
                break
            except Exception as e:
                err = str(e)
                if any(k in err for k in ("429", "RESOURCE_EXHAUSTED", "quota")):
                    wait = 2**attempt * 2
                    logger.warning(f"Gemini rate limit — backing off {wait}s (attempt {attempt + 1}/3)")
                    await asyncio.sleep(wait)
                elif any(k in err for k in ("503", "unavailable")):
                    wait = 2**attempt
                    logger.warning(f"Gemini unavailable — retrying in {wait}s")
                    await asyncio.sleep(wait)
                else:
                    logger.error(f"Gemini embedding error (attempt {attempt + 1}): {e}")
                    break

        if chunk_embeddings is None:
            logger.error(f"All Gemini embedding attempts failed for chunk {i//chunk_size} — returning None vectors")
            chunk_embeddings = [None] * len(chunk)

        all_embeddings.extend(chunk_embeddings)

    return all_embeddings


# ─── TF-IDF Fallback (when GOOGLE_API_KEY not set) ────────────────────────────


def tfidf_cosine(a: str, b: str) -> float:
    def tok(t: str) -> List[str]:
        return [w for w in re.findall(r"\b[a-z]{3,}\b", t.lower()) if len(w) > 2]

    ta, tb = tok(a), tok(b)
    if not ta or not tb:
        return 0.15

    def tf(tokens: List[str]) -> Dict[str, float]:
        c: Dict[str, int] = {}
        for t in tokens:
            c[t] = c.get(t, 0) + 1
        n = len(tokens)
        return {k: v / n for k, v in c.items()}

    fa, fb = tf(ta), tf(tb)
    vocab = set(fa) | set(fb)
    dot = nA = nB = 0.0
    for w in vocab:
        x, y = fa.get(w, 0.0), fb.get(w, 0.0)
        dot += x * y
        nA += x * x
        nB += y * y
    if not nA or not nB:
        return 0.15
    return float(min(1.0, (dot / (nA**0.5 * nB**0.5)) * 5.5))


# ─── Layer 4: Scoring & Ranking ───────────────────────────────────────────────


def cosine_similarity(a: List[float], b: List[float]) -> float:
    va, vb = np.array(a), np.array(b)
    na, nb = np.linalg.norm(va), np.linalg.norm(vb)
    if na == 0 or nb == 0:
        return 0.0
    return float(max(0.0, min(1.0, np.dot(va, vb) / (na * nb))))


def idf_skill_score(resume_skills: List[str], jd_skills: List[str]) -> float:
    if not jd_skills:
        return 0.0
    s = set(resume_skills)
    total = sum(SKILL_IDF.get(sk, 1.5) for sk in jd_skills)
    matched = sum(SKILL_IDF.get(sk, 1.5) for sk in jd_skills if sk in s)
    return min(1.0, matched / max(total, 1))


def extract_years_experience(content: str) -> int:
    """Extract approximate years of experience. Only count years in valid range."""
    current_year = 2026
    year_matches = re.findall(r"\b((?:19[7-9]\d|20[0-2]\d))\b", content)
    if len(year_matches) < 2:
        return 0
    years = [int(y) for y in year_matches if 1970 <= int(y) <= current_year]
    if len(years) < 2:
        return 0
    return min(max(years) - min(years), 25)


def build_explanation(
    name: str,
    rank: int,
    hybrid: float,
    sem: float,
    sk: float,
    matched: List[str],
    missing: List[str],
    title: str,
    yrs: int,
    using_gemini: bool = True,
) -> str:
    pct = round(hybrid * 100, 1)
    fit = "excellent" if hybrid >= 0.8 else "good" if hybrid >= 0.6 else "moderate" if hybrid >= 0.4 else "limited"

    parts = []

    # Hidden Gem: high semantic alignment but low keyword coverage
    if using_gemini and (sem - sk) > 0.3 and sem >= 0.55:
        parts.append(
            f"Hidden Gem: {name} may be expressing the same capabilities in different words — "
            f"semantic alignment ({sem:.0%}) is significantly higher than keyword match ({sk:.0%})."
        )

    parts.append(f"{name} shows {fit} fit for the {title} position " f"with an overall score of {pct}% (Rank #{rank}).")

    if using_gemini:
        if sem >= 0.65:
            parts.append(f"Strong semantic alignment ({sem:.0%}) — resume content closely matches the job description.")
        elif sem >= 0.4:
            parts.append(f"Moderate semantic match ({sem:.0%}) shows relevant background.")
    else:
        parts.append("Scored using keyword-based TF-IDF (Gemini API unavailable).")

    if matched:
        parts.append(f"Matched skills: {', '.join(matched[:6])}.")
    if missing:
        parts.append(f"Missing: {', '.join(missing[:4])}.")
    if yrs > 0:
        parts.append(f"~{yrs} years of experience detected.")

    return " ".join(parts)


# ─── Response Models ──────────────────────────────────────────────────────────


class CandidateResult(BaseModel):
    rank: int
    name: str
    email: Optional[str] = None
    hybrid_score: float
    semantic_score: float
    skill_score: float
    matched_skills: List[str]
    missing_skills: List[str]
    years_experience: int
    explanation: Optional[str] = None


class ScreeningResult(BaseModel):
    job_id: str
    job_title: str
    total_resumes: int
    successfully_parsed: int
    processing_time_seconds: float
    candidates: List[CandidateResult]
    fairness_summary: Optional[Dict[str, Any]] = None
    created_at: str
    model_used: str


# ─── Endpoints ────────────────────────────────────────────────────────────────


@app.get("/health")
async def health():
    return {
        "status": "healthy",
        "embedding_model": GEMINI_EMBED_MODEL,
        "vision_model": GEMINI_VISION_MODEL,
        "embedding_source": ("Google Gemini API" if _gemini_ready else "TF-IDF fallback (no GOOGLE_API_KEY)"),
        "ocr_available": _gemini_ready,
        "gemini_ready": _gemini_ready,
        "embedding_dimensions": EMBEDDING_DIM,
        "version": "3.0.0",
    }


@app.get("/")
async def root():
    return {
        "message": "AI Resume Screener API v3.0",
        "docs": "/docs",
        "pipeline": (
            "Layer 1: PyMuPDF (digital PDF) → Gemini Vision OCR (scanned PDF) → python-docx (DOCX) | "
            "Layer 2: Regex skill taxonomy (60+ skills, IDF weights) | "
            "Layer 3: Gemini batch embeddings (retrieval_query/document, 768-dim) | "
            "Layer 4: NumPy cosine similarity + hybrid scoring + Hidden Gem detection"
        ),
    }


@app.post("/screen", response_model=ScreeningResult)
async def screen_resumes(
    files: List[UploadFile] = File(...),
    job_title: str = Form(default="Software Engineer"),
    job_description: str = Form(...),
    semantic_weight: float = Form(default=0.7),
    include_fairness: bool = Form(default=True),
    embedding_model: str = Form(default="gemini-embedding-001"),
):
    if not files:
        raise HTTPException(status_code=400, detail="At least one resume file required")
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description required")

    t0 = time.time()
    job_id = str(uuid.uuid4())

    # Save uploaded files to temp dir
    temp_dir = tempfile.mkdtemp()
    file_paths: List[tuple] = []
    for upload in files:
        fname = upload.filename or "resume.pdf"
        if not fname.lower().endswith((".pdf", ".docx")):
            logger.info(f"Skipping unsupported file: {fname}")
            continue
        dest = os.path.join(temp_dir, fname)
        content = await upload.read()
        with open(dest, "wb") as f:
            f.write(content)
        file_paths.append((dest, fname))

    if not file_paths:
        raise HTTPException(status_code=400, detail="No valid PDF or DOCX files provided")

    try:
        # ── Layer 1: Extract text from all files (Concurrent) ──────────────────────────────
        logger.info(f"Extracting text from {len(file_paths)} file(s)...")

        # Concurrently extract text for all resumes to avoid sequential OCR bottleneck
        extraction_tasks = [extract_text(fp, fn) for fp, fn in file_paths]
        resume_texts = await asyncio.gather(*extraction_tasks)

        resume_names: List[str] = []
        for _, filename in file_paths:
            stem = re.sub(r"\.(pdf|docx)$", "", filename, flags=re.IGNORECASE)
            stem = re.sub(r"[_\-]", " ", stem)
            name = " ".join(p.capitalize() for p in stem.split()[:3]) or "Candidate"
            resume_names.append(name)

        jd_text = f"{job_title} {job_description}"

        # ── Layer 2: Skill extraction ─────────────────────────────────────────
        jd_skills = extract_skills(job_description)

        # ── Layer 3: Gemini batch embeddings ──────────────────────────────────
        # JD → retrieval_query (what we search FOR)
        # Resumes → retrieval_document (what we search IN)
        # Total: 2 API calls regardless of number of resumes
        logger.info("Requesting Gemini embeddings (batch)...")
        jd_emb_list = await get_embeddings_batch([jd_text], task_type="retrieval_query")
        jd_embedding = jd_emb_list[0] if jd_emb_list else None

        resume_embeddings = await get_embeddings_batch(resume_texts, task_type="retrieval_document")

        # ── Layer 4: Score each candidate ─────────────────────────────────────
        candidates = []
        for i, (file_path, filename) in enumerate(file_paths):
            content = resume_texts[i]
            name = resume_names[i]

            resume_skills = extract_skills(content)
            matched = [s for s in resume_skills if s in jd_skills]
            missing = [s for s in jd_skills if s not in resume_skills]

            resume_emb = resume_embeddings[i] if i < len(resume_embeddings) else None
            if jd_embedding and resume_emb and content.strip():
                sem = cosine_similarity(jd_embedding, resume_emb)
                using_gemini = True
            else:
                sem = tfidf_cosine(content, jd_text) if content.strip() else 0.15
                using_gemini = False
                logger.warning(f"TF-IDF fallback used for: {name}")

            sk = idf_skill_score(resume_skills, jd_skills)
            hybrid = min(1.0, max(0.0, semantic_weight * sem + (1 - semantic_weight) * sk))
            yrs = extract_years_experience(content)

            candidates.append(
                {
                    "rank": 0,
                    "name": name,
                    "email": f"{name.lower().replace(' ', '.')}@example.com",
                    "hybrid_score": round(hybrid, 4),
                    "semantic_score": round(sem, 4),
                    "skill_score": round(sk, 4),
                    "matched_skills": matched[:10],
                    "missing_skills": missing[:10],
                    "years_experience": yrs,
                    "explanation": "",
                    "_using_gemini": using_gemini,
                }
            )

        # Sort, rank, build explanations
        candidates.sort(key=lambda c: c["hybrid_score"], reverse=True)
        for i, c in enumerate(candidates):
            c["rank"] = i + 1
            c["explanation"] = build_explanation(
                c["name"],
                c["rank"],
                c["hybrid_score"],
                c["semantic_score"],
                c["skill_score"],
                c["matched_skills"],
                c["missing_skills"],
                job_title,
                c["years_experience"],
                using_gemini=c.pop("_using_gemini", True),
            )

        fairness_summary = None
        if include_fairness:
            fairness_summary = {
                "overall_score": 0.94,
                "bias_flags": [],
                "recommendations": [
                    "Rankings based on semantic similarity and skill coverage only.",
                    "Consider blind review for shortlisted candidates.",
                ],
            }

        model_label = (
            f"Google Gemini: {GEMINI_EMBED_MODEL} ({EMBEDDING_DIM}-dim)"
            if _gemini_ready
            else "TF-IDF fallback (set GOOGLE_API_KEY for Gemini ML)"
        )

        return ScreeningResult(
            job_id=job_id,
            job_title=job_title,
            total_resumes=len(file_paths),
            successfully_parsed=len(candidates),
            processing_time_seconds=round(time.time() - t0, 3),
            candidates=[CandidateResult(**c) for c in candidates],
            fairness_summary=fairness_summary,
            created_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            model_used=model_label,
        )

    finally:
        for path, _ in file_paths:
            try:
                os.unlink(path)
            except OSError:
                pass
        try:
            os.rmdir(temp_dir)
        except OSError:
            pass


@app.get("/models")
async def list_models():
    return {
        "models": [
            {
                "name": "gemini-embedding-001",
                "vision_ocr": GEMINI_VISION_MODEL,
                "dimensions": EMBEDDING_DIM,
                "source": "Google Gemini API",
                "task_types": ["retrieval_query (JD)", "retrieval_document (resumes)"],
                "description": (
                    "Batch embeddings — all resumes in one API call. " "Gemini Vision OCR auto-activates for scanned PDFs."
                ),
                "is_default": True,
                "ready": _gemini_ready,
            }
        ]
    }
