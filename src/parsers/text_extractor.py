"""Text extraction from PDF and DOCX files."""

import logging
import traceback
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


class TextExtractor:
    """Handles text extraction from various document formats."""
    
    def __init__(self):
        """Initialize text extractor."""
        self._pdf_extractors = []
        self._docx_extractors = []
        
        # Try to import PDF libraries
        try:
            import fitz  # PyMuPDF
            self._pdf_extractors.append(self._extract_with_pymupdf)
            logger.debug("PyMuPDF available for PDF extraction")
        except ImportError:
            logger.warning("PyMuPDF not available")
        
        try:
            import pdfplumber
            self._pdf_extractors.append(self._extract_with_pdfplumber)
            logger.debug("pdfplumber available for PDF extraction")
        except ImportError:
            logger.warning("pdfplumber not available")
        
        # Try to import DOCX library
        try:
            import docx
            self._docx_extractors.append(self._extract_with_python_docx)
            logger.debug("python-docx available for DOCX extraction")
        except ImportError:
            logger.warning("python-docx not available")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                return self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            logger.debug(f"Full traceback: {traceback.format_exc()}")
            return f"Error extracting text: {str(e)}"
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using available libraries."""
        if not self._pdf_extractors:
            raise ValueError("No PDF extraction libraries available. Install PyMuPDF or pdfplumber.")
        
        # Try each PDF extractor until one works
        for extractor in self._pdf_extractors:
            try:
                text = extractor(file_path)
                if text and text.strip():
                    return text
            except Exception as e:
                logger.debug(f"PDF extractor failed: {str(e)}")
                continue
        
        # If all extractors fail, return error message
        return "Failed to extract text from PDF file"
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX using available libraries."""
        if not self._docx_extractors:
            raise ValueError("No DOCX extraction libraries available. Install python-docx.")
        
        # Try each DOCX extractor until one works
        for extractor in self._docx_extractors:
            try:
                text = extractor(file_path)
                if text and text.strip():
                    return text
            except Exception as e:
                logger.debug(f"DOCX extractor failed: {str(e)}")
                continue
        
        # If all extractors fail, return error message
        return "Failed to extract text from DOCX file"
    
    def _extract_with_pymupdf(self, file_path: str) -> str:
        """Extract text using PyMuPDF (fitz)."""
        import fitz
        
        text_content = []
        total_chars = 0
        total_images = 0
        
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                # Count characters and images for OCR detection
                total_chars += len(text.strip())
                image_list = page.get_images()
                total_images += len(image_list)
                
                if text.strip():
                    text_content.append(text)
        
        extracted_text = "\\n".join(text_content)
        
        # Check if this might be a scanned PDF (low text-to-image ratio)
        if total_images > 0 and total_chars < (total_images * 50):
            ocr_warning = ("\\n\\n[OCR WARNING: This appears to be a scanned PDF with limited extractable text. "
                          "For better results, consider using OCR software to convert the document to searchable text.]\\n")
            extracted_text = ocr_warning + extracted_text
            logger.warning(f"Potential scanned PDF detected: {file_path} - {total_chars} chars, {total_images} images")
        
        return extracted_text
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber (fallback for complex layouts)."""
        import pdfplumber
        
        text_content = []
        total_chars = 0
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Try to extract text
                text = page.extract_text()
                
                if text and text.strip():
                    text_content.append(text)
                    total_chars += len(text.strip())
                
                # Also try to extract text from tables if regular extraction fails
                if not text or len(text.strip()) < 50:
                    try:
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    row_text = " | ".join([cell or "" for cell in row])
                                    if row_text.strip():
                                        text_content.append(row_text)
                                        total_chars += len(row_text.strip())
                    except Exception as e:
                        logger.debug(f"Table extraction failed on page: {str(e)}")
        
        extracted_text = "\\n".join(text_content)
        
        # Check for potential OCR needs with pdfplumber
        if total_chars < 100:
            ocr_warning = ("\\n\\n[OCR WARNING: Very little text was extracted from this PDF. "
                          "This may be a scanned document that requires OCR processing for better text extraction.]\\n")
            extracted_text = ocr_warning + extracted_text
            logger.warning(f"Low text extraction with pdfplumber: {file_path} - only {total_chars} characters")
        
        return extracted_text
    
    def _extract_with_python_docx(self, file_path: str) -> str:
        """
        Extract text using python-docx with structured extraction.
        
        Handles tables, headers, and formatting while preserving section structure.
        """
        import docx
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt
        
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract headers and footers first
        headers_footers = self._extract_headers_footers(doc)
        if headers_footers:
            text_content.extend(headers_footers)
            text_content.append("=" * 50)  # Separator
        
        # Process document body with structure preservation
        current_section = []
        last_style = None
        
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                paragraph = None
                # Find the corresponding paragraph object
                for p in doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                
                if paragraph and paragraph.text.strip():
                    formatted_text = self._format_paragraph_text(paragraph)
                    
                    # Detect section breaks based on style and formatting
                    if self._is_section_header(paragraph):
                        # Save previous section if it exists
                        if current_section:
                            text_content.extend(current_section)
                            text_content.append("")  # Add spacing between sections
                            current_section = []
                        
                        # Add section header with emphasis
                        text_content.append(f"\n## {formatted_text}")
                        last_style = 'header'
                    else:
                        current_section.append(formatted_text)
                        last_style = 'paragraph'
            
            # Handle tables
            elif element.tag.endswith('tbl'):
                # Find the corresponding table object
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table:
                    table_text = self._extract_table_structured(table)
                    if table_text:
                        # Save current section before table
                        if current_section:
                            text_content.extend(current_section)
                            current_section = []
                        
                        text_content.append("\n[TABLE]")
                        text_content.extend(table_text)
                        text_content.append("[/TABLE]\n")
        
        # Add any remaining content
        if current_section:
            text_content.extend(current_section)
        
        return "\n".join(text_content)
    
    def _extract_headers_footers(self, doc) -> list:
        """Extract text from headers and footers."""
        headers_footers = []
        
        try:
            # Extract headers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            headers_footers.append(f"[HEADER] {paragraph.text.strip()}")
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            headers_footers.append(f"[FOOTER] {paragraph.text.strip()}")
        except Exception as e:
            logger.debug(f"Error extracting headers/footers: {str(e)}")
        
        return headers_footers
    
    def _format_paragraph_text(self, paragraph) -> str:
        """Format paragraph text preserving important formatting cues."""
        text = paragraph.text.strip()
        
        if not text:
            return ""
        
        # Check for bold/italic formatting in runs
        formatted_parts = []
        for run in paragraph.runs:
            run_text = run.text
            if run_text.strip():
                # Preserve bold formatting with markdown-style emphasis
                if run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                formatted_parts.append(run_text)
        
        # If we have formatted parts, use them; otherwise use plain text
        if formatted_parts:
            formatted_text = "".join(formatted_parts)
        else:
            formatted_text = text
        
        return formatted_text
    
    def _is_section_header(self, paragraph) -> bool:
        """Determine if a paragraph is likely a section header."""
        text = paragraph.text.strip()
        
        if not text:
            return False
        
        # Check style name for header indicators
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if any(header_word in style_name for header_word in ['heading', 'title', 'header']):
            return True
        
        # Check formatting characteristics
        if paragraph.runs:
            first_run = paragraph.runs[0]
            # Bold text that's short might be a header
            if first_run.bold and len(text) < 100:
                return True
            
            # Check font size (headers often have larger fonts)
            try:
                if first_run.font.size and first_run.font.size >= Pt(14):
                    return True
            except Exception:
                pass  # Font size might not be available
        
        # Check for common header patterns
        header_patterns = [
            text.isupper() and len(text) < 50,  # ALL CAPS short text
            text.endswith(':') and len(text) < 80,  # Text ending with colon
            len(text.split()) <= 5 and not text.endswith('.'),  # Short phrases without periods
        ]
        
        return any(header_patterns)
    
    def _extract_table_structured(self, table) -> list:
        """Extract table content with structure preservation."""
        table_content = []
        
        try:
            # Check if first row might be headers
            if table.rows:
                first_row = table.rows[0]
                first_row_cells = [cell.text.strip() for cell in first_row.cells]
                
                # If first row has bold formatting or short text, treat as headers
                is_header_row = False
                if first_row.cells:
                    for cell in first_row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.bold:
                                    is_header_row = True
                                    break
                            if is_header_row:
                                break
                        if is_header_row:
                            break
                
                # Process all rows
                for i, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                        else:
                            row_cells.append("")
                    
                    if any(cell for cell in row_cells):  # Only add non-empty rows
                        if i == 0 and is_header_row:
                            # Format header row
                            table_content.append("| " + " | ".join(row_cells) + " |")
                            table_content.append("|" + "|".join([" --- " for _ in row_cells]) + "|")
                        else:
                            # Format data row
                            table_content.append("| " + " | ".join(row_cells) + " |")
        
        except Exception as e:
            logger.debug(f"Error extracting table structure: {str(e)}")
            # Fallback to simple extraction
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_content.append(" | ".join(row_text))
        
        return table_content
    
    def is_scanned_pdf(self, file_path: str) -> bool:
        """
        Detect if a PDF is likely scanned and may need OCR.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            True if the PDF appears to be scanned, False otherwise
        """
        path = Path(file_path)
        
        if path.suffix.lower() != '.pdf':
            return False
        
        try:
            import fitz
            
            with fitz.open(file_path) as doc:
                total_chars = 0
                total_images = 0
                
                # Check first few pages for efficiency
                pages_to_check = min(3, len(doc))
                
                for page_num in range(pages_to_check):
                    page = doc[page_num]
                    text = page.get_text()
                    total_chars += len(text.strip())
                    
                    # Count images on the page
                    image_list = page.get_images()
                    total_images += len(image_list)
                
                # Heuristics for scanned PDF detection:
                # 1. Very few characters relative to number of images
                # 2. Very low character count overall
                # 3. High image-to-text ratio
                
                if total_chars < 50:  # Very little text
                    return True
                
                if total_images > 0 and total_chars < (total_images * 100):  # Low text-to-image ratio
                    return True
                
                # Check character density per page
                avg_chars_per_page = total_chars / pages_to_check if pages_to_check > 0 else 0
                if avg_chars_per_page < 200:  # Very sparse text
                    return True
                
                return False
                
        except Exception as e:
            logger.debug(f"Error checking if PDF is scanned: {str(e)}")
            return False
    
    def get_document_info(self, file_path: str) -> dict:
        """
        Get metadata information about the document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with document metadata
        """
        path = Path(file_path)
        
        info = {
            'file_name': path.name,
            'file_size': path.stat().st_size,
            'file_extension': path.suffix.lower(),
            'extractors_available': {
                'pdf': len(self._pdf_extractors),
                'docx': len(self._docx_extractors)
            }
        }
        
        # Try to get additional metadata for PDFs
        if path.suffix.lower() == '.pdf' and self._pdf_extractors:
            try:
                import fitz
                with fitz.open(file_path) as doc:
                    info.update({
                        'page_count': len(doc),
                        'title': doc.metadata.get('title', ''),
                        'author': doc.metadata.get('author', ''),
                        'creator': doc.metadata.get('creator', ''),
                        'is_scanned': self.is_scanned_pdf(file_path)
                    })
            except Exception as e:
                logger.debug(f"Could not extract PDF metadata: {str(e)}")
        
        return info